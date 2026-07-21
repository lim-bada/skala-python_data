"""
프로그램명: 창의적 개인 실습 보고서 Word 변환기
작성자: 임해안
작성일: 2026-07-21

목적:
    - REPORT.md의 제목, 문단, 목록, 표와 이미지를 편집 가능한 Word 문서로 변환한다.
    - 캡처 이미지를 교체한 뒤 같은 명령으로 보고서를 다시 생성할 수 있게 한다.
"""

from pathlib import Path

import markdown
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


REPORT_DIR = Path(__file__).resolve().parent
MARKDOWN_PATH = REPORT_DIR / "REPORT.md"
WORD_PATH = REPORT_DIR / "Advanced_창의적실습_보고서.docx"
BODY_FONT = "맑은 고딕"
CODE_FONT = "Consolas"


def set_east_asia_font(run: object, font_name: str) -> None:
    """한글이 지정한 글꼴로 표시되도록 Word의 동아시아 글꼴을 설정한다."""
    run.font.name = font_name  # type: ignore[attr-defined]
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)  # type: ignore[attr-defined]


def set_cell_shading(cell: object, fill: str) -> None:
    """표 머리글이나 강조 셀에 배경색을 적용한다."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)  # type: ignore[attr-defined]


def set_paragraph_shading(paragraph: object, fill: str) -> None:
    """인용문과 코드 블록 문단에 배경색을 적용한다."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    paragraph._p.get_or_add_pPr().append(shading)  # type: ignore[attr-defined]


def prevent_row_split(row: object) -> None:
    """표의 한 행이 페이지 경계에서 둘로 나뉘지 않도록 설정한다."""
    cant_split = OxmlElement("w:cantSplit")
    row._tr.get_or_add_trPr().append(cant_split)  # type: ignore[attr-defined]


def add_page_number(paragraph: object) -> None:
    """바닥글 문단에 Word PAGE 필드를 삽입한다."""
    paragraph.add_run("창의적 개인 실습 보고서  |  ")  # type: ignore[attr-defined]
    run = paragraph.add_run()  # type: ignore[attr-defined]
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def configure_document(document: Document) -> None:
    """A4 용지, 한글 글꼴, 제목 스타일과 바닥글을 설정한다."""
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(16)
    section.bottom_margin = Mm(16)
    section.left_margin = Mm(17)
    section.right_margin = Mm(17)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.35

    for style_name, size, color in (
        ("Title", 24, "183153"),
        ("Heading 1", 16, "183153"),
        ("Heading 2", 12.5, "1F4B99"),
    ):
        style = styles[style_name]
        style.font.name = BODY_FONT
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.paragraph_format.keep_with_next = True

    styles["Title"].font.bold = True
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.bold = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer)
    for run in footer.runs:
        set_east_asia_font(run, BODY_FONT)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 116, 139)

    properties = document.core_properties
    properties.title = "웹 로그 기반 서비스 장애 조기탐지 및 원인 분석 시스템"
    properties.author = "광주_3반_임해안"
    properties.subject = "창의적 개인 실습 보고서"


def add_inline_content(paragraph: object, node: object) -> None:
    """HTML 인라인 태그를 Word run으로 변환하며 굵기와 코드 스타일을 유지한다."""
    if isinstance(node, NavigableString):
        text = str(node)
        if text:
            run = paragraph.add_run(text)  # type: ignore[attr-defined]
            set_east_asia_font(run, BODY_FONT)
        return
    if not isinstance(node, Tag):
        return

    if node.name == "br":
        paragraph.add_run().add_break()  # type: ignore[attr-defined]
        return

    text = node.get_text()
    if node.name in {"strong", "b", "em", "i", "code", "a"}:
        run = paragraph.add_run(text)  # type: ignore[attr-defined]
        set_east_asia_font(run, BODY_FONT)
        if node.name in {"strong", "b"}:
            run.bold = True
            run.font.color.rgb = RGBColor(24, 49, 83)
        if node.name in {"em", "i"}:
            run.italic = True
        if node.name == "code":
            set_east_asia_font(run, CODE_FONT)
            run.font.color.rgb = RGBColor(180, 35, 24)
            run.font.size = Pt(9)
        if node.name == "a":
            run.font.color.rgb = RGBColor(37, 99, 235)
            run.underline = True
        return

    for child in node.children:
        add_inline_content(paragraph, child)


def add_text_paragraph(
    document: Document,
    element: Tag,
    *,
    style: str | None = None,
) -> object:
    """HTML 문단이나 목록 항목을 Word 문단으로 변환한다."""
    paragraph = document.add_paragraph(style=style)
    for child in element.children:
        add_inline_content(paragraph, child)
    return paragraph


def add_image(document: Document, image: Tag) -> None:
    """Markdown 상대 경로 이미지를 본문 너비에 맞춰 삽입한다."""
    source = image.get("src")
    if not source:
        return
    image_path = REPORT_DIR / source
    if not image_path.is_file():
        paragraph = document.add_paragraph()
        paragraph.add_run(f"[이미지 없음: {source}]").font.color.rgb = RGBColor(
            185, 28, 28
        )
        return

    document.add_picture(str(image_path), width=Inches(6.45))
    paragraph = document.paragraphs[-1]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.keep_together = True


def add_table(document: Document, html_table: Tag) -> None:
    """Markdown 표를 편집 가능한 Word 표로 변환한다."""
    html_rows = html_table.find_all("tr")
    if not html_rows:
        return
    column_count = max(len(row.find_all(["th", "td"])) for row in html_rows)
    table = document.add_table(rows=len(html_rows), cols=column_count)
    table.style = "Table Grid"
    table.autofit = True

    for row_index, html_row in enumerate(html_rows):
        word_row = table.rows[row_index]
        prevent_row_split(word_row)
        for column_index, html_cell in enumerate(html_row.find_all(["th", "td"])):
            cell = word_row.cells[column_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            for child in html_cell.children:
                add_inline_content(paragraph, child)
            for run in paragraph.runs:
                run.font.size = Pt(8.5)
                if row_index == 0 or html_cell.name == "th":
                    run.bold = True
            if row_index == 0 or html_cell.name == "th":
                set_cell_shading(cell, "EAF1FB")
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def add_code_block(document: Document, element: Tag) -> None:
    """코드 블록을 고정폭 글꼴과 배경색이 있는 Word 문단으로 변환한다."""
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.left_indent = Mm(3)
    paragraph.paragraph_format.right_indent = Mm(3)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(7)
    set_paragraph_shading(paragraph, "172033")
    run = paragraph.add_run(element.get_text())
    set_east_asia_font(run, CODE_FONT)
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(248, 250, 252)


def convert_markdown_to_word() -> Path:
    """REPORT.md를 읽어 구조와 이미지를 유지한 Word 문서를 생성한다."""
    source = MARKDOWN_PATH.read_text(encoding="utf-8")
    html = markdown.markdown(
        source,
        extensions=["tables", "fenced_code", "sane_lists"],
    )
    soup = BeautifulSoup(html, "html.parser")
    document = Document()
    configure_document(document)

    for element in soup.children:
        if not isinstance(element, Tag):
            continue
        if element.name == "h1":
            paragraph = add_text_paragraph(document, element, style="Title")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif element.name == "h2":
            add_text_paragraph(document, element, style="Heading 1")
        elif element.name == "h3":
            add_text_paragraph(document, element, style="Heading 2")
        elif element.name == "p":
            image = element.find("img")
            if image:
                add_image(document, image)
            else:
                add_text_paragraph(document, element)
        elif element.name == "blockquote":
            paragraph = document.add_paragraph()
            for child in element.children:
                add_inline_content(paragraph, child)
            paragraph.paragraph_format.left_indent = Mm(5)
            paragraph.paragraph_format.right_indent = Mm(5)
            set_paragraph_shading(paragraph, "EEF4FF")
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(36, 64, 100)
        elif element.name in {"ul", "ol"}:
            list_style = "List Bullet" if element.name == "ul" else "List Number"
            for item in element.find_all("li", recursive=False):
                add_text_paragraph(document, item, style=list_style)
        elif element.name == "table":
            add_table(document, element)
        elif element.name == "pre":
            add_code_block(document, element)
        elif element.name == "hr":
            document.add_paragraph()

    document.save(WORD_PATH)
    return WORD_PATH


def main() -> None:
    """Word 보고서를 생성하고 저장 경로와 파일 크기를 출력한다."""
    output_path = convert_markdown_to_word()
    print(f"Word 생성 완료: {output_path}")
    print(f"Word 파일 크기: {output_path.stat().st_size / 1024:,.1f} KB")


if __name__ == "__main__":
    main()
